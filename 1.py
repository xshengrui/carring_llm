eval.py:
import json
from pathlib import Path
import argparse
from transformers import AutoTokenizer
from vllm import LLM, SamplingParams

JUDGE_PROMPT = """
你是一名专业的医学护理知识问答评测员。你的任务是比较模型的回答与标准答案之间的知识要点是否一致。  
请依据以下标准进行评分：

评分标准（总分0-2分）：
- 0分：回答错误
- 1分：部分回答正确
- 2分：完全回答正确

【问题】{question}  
【标准答案】{answer}  
【模型回答】{model_response}  

评分（仅输出一个阿拉伯数字）：
"""

def get_parser():
    parser = argparse.ArgumentParser(description="Evaluate model outputs.")
    parser.add_argument('--input_file', type=str, required=True)
    parser.add_argument('--output_file', type=str, required=True)
    parser.add_argument('--model_path', type=str, default="Qwen/Qwen3-32B")
    parser.add_argument('--gpus', type=int, default=4, help='Number of GPUs to use')
    return parser

def init_data(file_path):
    data = []
    with open(file_path, 'r', encoding='utf-8') as f:
        for line in f:
            if line.strip():
                data.append(json.loads(line))
    for item in data:
        item['prompt_judge_score'] = []
    return data


def build_prompt(question, answer, model_response):
    user_prompt =  JUDGE_PROMPT.format(
        question=question,
        answer=answer,
        model_response=model_response
    )
    messages = [{"role": "user", "content": user_prompt}]
    text = tokenizer.apply_chat_template(
        messages,
        tokenize=False,
        add_generation_prompt=True,
        enable_thinking=False, 
    )
    return text

def eval_prompt_judge(input_path, output_path):
    data = init_data(input_path)
    prompts = [build_prompt(item["question"], item["answer"], item["model_response"]) for item in data]

    for i in range(3):
        outputs = llm.generate(prompts, sampling_params)
        for idx, output in enumerate(outputs):
            text = output.outputs[0].text.strip()
            try:
                score = int(''.join(filter(str.isdigit, text)))
                if 0 <= score <= 2:
                    data[idx]['prompt_judge_score'].append(score)
            except:
                pass

    for item in data:
        if item['prompt_judge_score']:
            item['prompt_judge_score'] = sum(item['prompt_judge_score']) / len(item['prompt_judge_score'])
        else:
            item['prompt_judge_score'] = 0

    with open(output_path, 'w', encoding='utf-8') as fout:
        for item in data:
            fout.write(json.dumps(item, ensure_ascii=False) + '\n')

def calc_and_print_avg_score(output_path: str):
    scores = []
    with open(output_path, 'r', encoding='utf-8') as fin:
        for line in fin:
            data = json.loads(line)
            score = data.get('prompt_judge_score', None)
            if score is not None and score != -1:
                scores.append(score * 50)
    if scores:
        avg = sum(scores) / len(scores)
        print(f"平均分: {avg:.2f} (0-100分制, 样本数: {len(scores)})")
    else:
        print("未找到有效分数，无法计算平均分。")

def main():
    parser = get_parser()
    args = parser.parse_args()
    input_path = args.input_file
    output_path = args.output_file
    model_path = args.model_path
    gpus = args.gpus

    import os
    if not os.path.exists(input_path):
        print("input file 未找到")
        return 0
    
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    global llm
    global tokenizer
    tokenizer = AutoTokenizer.from_pretrained(model_path)
    global sampling_params 
    sampling_params = SamplingParams(
        temperature=0.6, top_p=0.95, top_k=20, max_tokens=64
    )

    llm = LLM(
        model=model_path,
        tensor_parallel_size=gpus,
        max_num_seqs=16
    )
    
    eval_prompt_judge(input_path, output_path)
    calc_and_print_avg_score(output_path)

if __name__ == "__main__":
    main()


inference_x.py:
import argparse
import json
from tqdm import tqdm
from transformers import AutoTokenizer
from vllm import LLM, SamplingParams

def parse_args():
    parser = argparse.ArgumentParser(description="推理脚本")
    parser.add_argument("--input_file", type=str, required=True, help="输入 JSONL 文件路径")
    parser.add_argument("--output_file", type=str, required=True, help="输出 JSONL 文件路径")
    parser.add_argument("--model_path", type=str, required=True, help="模型路径")
    parser.add_argument("--gpus", type=int, default=1, help="使用的 GPU 数量")
    return parser.parse_args()

####
# 这是一个示例infrence.py脚本，用于输入指定格式的 JSONL 文件，调用 Qwen3-8B 模型进行推理，并将结果保存到输出文件中。
# 大家需要根据自己的技术方案，如检索增强、多智能体系统等，自行实现推理逻辑。该脚本仅用于给出遵循输入输出格式的简单示例。
# 建议修改成batch inference的形式
####

# --- 提示词构建函数 ---
def build_prompt_default(question, tokenizer):
    messages = [{"role": "user", "content": question}]
    prompt = tokenizer.apply_chat_template(
        messages,
        tokenize=False,
        add_generation_prompt=True,
        enable_thinking=False,
    )
    return prompt

def build_prompt_strategy_a(question, tokenizer):
    """策略 A：明确角色与指令"""
    user_prompt = (
        "你是一名专业的医疗护理专家。请根据你的专业知识，准确、详细地回答以下护理问题。\n"
        f"【问题】{question}"
    )
    messages = [{"role": "user", "content": user_prompt}]
    text = tokenizer.apply_chat_template(
        messages,
        tokenize=False,
        add_generation_prompt=True,
        enable_thinking=False,
    )
    return text

def build_prompt_strategy_b(question, tokenizer):
    """策略 B：引入 CoT (Chain-of-Thought) 思维链"""
    user_prompt = (
        "你是一名专业的医疗护理专家。请思考以下问题，并逐步推理，最终给出准确、详细的护理回答。\n"
        f"【问题】{question}\n"
        "【思考过程】"
    )
    messages = [{"role": "user", "content": user_prompt}]
    text = tokenizer.apply_chat_template(
        messages,
        tokenize=False,
        add_generation_prompt=True,
        enable_thinking=True, # Qwen模型可能支持enable_thinking来激活CoT
    )
    return text

def build_prompt_strategy_c(question, tokenizer):
    """策略 C：强调简洁与关键信息"""
    user_prompt = (
        "请你作为一名经验丰富的护理人员，针对以下问题，提供简洁但包含所有关键信息的回答。不要包含无关的细节。\n"
        f"【问题】{question}"
    )
    messages = [{"role": "user", "content": user_prompt}]
    text = tokenizer.apply_chat_template(
        messages,
        tokenize=False,
        add_generation_prompt=True,
        enable_thinking=False,
    )
    return text

# 假设有以下示例数据（通常需要从实际数据集中选取）
# 注意：在实际使用时，这个示例数据应该来自你的验证集或训练集，并且是高质量的。
example_qa_pair = {
    "question": "Ⅱ型呼衰病人应用高流量湿化氧疗时, 氧流量应设置为多少?",
    "answer": "氧流量在30L/min时,病人PaCO2有较大改善,建议最低流量应设置为 30L/min。"
}

def build_prompt_strategy_d(question, tokenizer, example_qa=example_qa_pair):
    """策略 D：结合示例 (Few-shot prompting)"""
    user_prompt = (
        "以下是一个护理问题及其专业回答示例：\n"
        f"【示例问题】{example_qa['question']}\n"
        f"【示例回答】{example_qa['answer']}\n\n"
        "请参考上述示例的风格和专业程度，准确回答以下新的护理问题。\n"
        f"【问题】{question}"
    )
    messages = [{"role": "user", "content": user_prompt}]
    text = tokenizer.apply_chat_template(
        messages,
        tokenize=False,
        add_generation_prompt=True,
        enable_thinking=False,
    )
    return text

def build_prompt_strategy_e(question, tokenizer):
    """策略 E：结合 CoT 与专家角色设定 (强化版 CoT)"""
    user_prompt = (
        "你是一名资深的医疗护理专家。请你针对以下护理问题，首先进行详细的专业思考，包括但不限于问题的关键点、可能涉及的护理原则、相关医学知识等。然后，基于你的思考，给出全面、准确、专业的最终回答。\n"
        f"【问题】{question}\n"
        "【专业思考过程】" # 更明确地引导专业思考
    )
    messages = [{"role": "user", "content": user_prompt}]
    text = tokenizer.apply_chat_template(
        messages,
        tokenize=False,
        add_generation_prompt=True,
        enable_thinking=True, # 启用CoT
    )
    return text

def build_prompt_strategy_f(question, tokenizer):
    """策略 F：多步骤思考与结构化输出 (ToT 简化版)"""
    user_prompt = (
        "你是一名严谨的医疗护理顾问。请严格按照以下步骤思考并回答护理问题：\n"
        "1. **识别核心问题**：明确用户问题的核心诉求和关键信息。\n"
        "2. **关联知识点**：思考与核心问题相关的护理学、医学理论、指南或临床实践经验。\n"
        "3. **制定初步回答**：根据关联知识点，构思一个初步的、全面的回答草稿。\n"
        "4. **审查与优化**：检查初步回答的准确性、完整性和专业性，并进行必要的修正和完善。\n"
        "最后，以清晰、条理分明的格式给出最终回答。\n"
        f"【问题】{question}\n"
        "【思考与回答】" # 引导模型按照步骤输出
    )
    messages = [{"role": "user", "content": user_prompt}]
    text = tokenizer.apply_chat_template(
        messages,
        tokenize=False,
        add_generation_prompt=True,
        enable_thinking=True, # 启用CoT
    )
    return text

def build_prompt_strategy_g(question, tokenizer):
    """策略 G：强调准确性与安全性 (医疗领域特化)"""
    user_prompt = (
        "你是一名专业的医疗护理人工智能助理，你的首要任务是提供准确、安全、负责任的护理知识。请在回答以下问题时，务必做到信息正确无误，避免任何可能误导或造成风险的表述。\n"
        "请先详细思考，确保理解问题的所有方面，并调动所有相关专业知识。如果信息存在不确定性或需要更具体的临床判断，请在回答中明确指出。\n"
        f"【问题】{question}\n"
        "【思考与回答】"
    )
    messages = [{"role": "user", "content": user_prompt}]
    text = tokenizer.apply_chat_template(
        messages,
        tokenize=False,
        add_generation_prompt=True,
        enable_thinking=True, # 启用CoT
    )
    return text
# --- 提示词构建函数结束 ---


def main():
    args = parse_args()

    # 加载分词器
    tokenizer = AutoTokenizer.from_pretrained(args.model_path)
    # 初始化LLM模型，设置GPU并行数量
    llm = LLM(model=args.model_path, tensor_parallel_size=args.gpus)

    # 设置生成参数
    sampling_params = SamplingParams(
        temperature=0.6,    # 控制生成文本的随机性
        top_p=0.95,        # 控制累积概率阈值
        top_k=20,          # 控制每次选择的候选词数量
        max_tokens=32768   # 生成文本的最大长度
    )

    # 读取输入文件
    with open(args.input_file, 'r', encoding='utf-8') as f:
        input_data = [json.loads(line) for line in f]  # 将每行JSONL数据解析为Python对象

    prompts = []
    # 遍历处理每个输入样本，构建所有prompt
    for item in input_data:
        question = item.get("question", "")
        
        # 选择你想要测试的提示词策略
        # 默认策略
        # prompts.append(build_prompt_default(question, tokenizer)) 
        
        # 策略 A：明确角色与指令
        # prompts.append(build_prompt_strategy_a(question, tokenizer))
        
        # 策略 B：引入 CoT 思维链
        # prompts.append(build_prompt_strategy_b(question, tokenizer))
        
        # 策略 C：强调简洁与关键信息
        # prompts.append(build_prompt_strategy_c(question, tokenizer))
        
        # 策略 D：结合示例 (Few-shot prompting) - 需要 example_qa_pair
        # prompts.append(build_prompt_strategy_d(question, tokenizer))
        
        # 策略 E：结合 CoT 与专家角色设定 (强化版 CoT)
        prompts.append(build_prompt_strategy_e(question, tokenizer))
        
        # 策略 F：多步骤思考与结构化输出 (ToT 简化版)
        # prompts.append(build_prompt_strategy_f(question, tokenizer))

    print(f"Total prompts to generate: {len(prompts)}")

    # 批量生成
    outputs = llm.generate(prompts, sampling_params)

    # 将vLLM的输出映射回原始的input_data
    # 由于prompt可能重复，这里需要更稳健的映射方式
    # outputs是按照输入prompts列表的顺序返回的
    
    # 存储输出结果的列表
    processed_outputs = []

    for i, item in enumerate(input_data):
        qid = item.get("id")
        question = item.get("question", "")
        answer = item.get("answer", "")

        # 确保输出结果与原始输入对应
        # 注意：这里假设outputs的顺序与prompts的顺序一致
        # vLLM默认会保持请求的顺序，所以可以通过索引直接对应
        generated_text = outputs[i].outputs[0].text.strip()
        
        output_item = {
            "id": qid,
            "question": question,
            "answer": answer,
            "model_response": generated_text
        }
        processed_outputs.append(output_item)

    # 将结果写入输出文件
    with open(args.output_file, 'w', encoding='utf-8') as f:
        for item in processed_outputs:
            f.write(json.dumps(item, ensure_ascii=False) + '\n')

# 程序入口点
if __name__ == "__main__":
    main()

RAG.py:
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
'''
python RAG.py \
    --input_file /inspire/hdd/project/socialscience/xialingying041-summer-041/project/data/val300.jsonl \
    --output_file ./results/rag/rag.jsonl \
    --model_path ./final_model \
    --gpus 1 \
    --batch_size 64 \
    --gpu_memory_utilization 0.95 \
    --max_num_batched_tokens 16384
'''
import argparse
import json
from tqdm import tqdm
from transformers import AutoTokenizer
from vllm import LLM, SamplingParams
import os
import math

# --- LangChain RAG 组件导入 ---
from langchain_community.document_loaders import DirectoryLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceBgeEmbeddings # 或 HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document # 用于类型提示
from typing import List

# --- RAG 数据相关配置 ---
RAG_DATA_PATH = "/inspire/hdd/project/socialscience/xialingying041-summer-041/pdf/output/extracted_texts"
RAG_EMBEDDING_MODEL_NAME = "/inspire/hdd/project/socialscience/xialingying041-summer-041/models/models--BAAI--bge-small-zh-v1.5" # 或其他合适的中文嵌入模型
RAG_CHUNK_SIZE = 500  # 文本分块大小
RAG_CHUNK_OVERLAP = 50 # 文本分块重叠大小
RAG_SEARCH_K = 3      # 检索最相似的文本块数量

# # --- 优化后的系统级 Prompt (保持原样) ---
# SYSTEM_INFERENCE_PROMPT = """你是一名专业的医学护理专家，你的任务是针对用户提出的护理问题，提供直接、准确、全面、专业的回答。你的回答应该基于权威的护理知识，并力求与最佳实践指南保持一致。
# 请确保你的回答：
# 1. 知识点准确无误。
# 2. 内容全面，涵盖问题涉及的关键要点。
# 3. 语言简洁明了，避免冗余信息。
# 4. 严格围绕问题，不进行不必要的扩展或发散。
# 5. 采用专业的医学护理术语和表达。
# 6. 如果问题涉及具体操作或流程，请尽可能详尽描述关键步骤。
# 7. 如果信息不确定或超出你的知识范围，请避免臆造，可以指出信息有限。
# """

# --- RAG 增强的用户 Prompt 模板 ---
# 在原始问题前加入检索到的参考资料
RAG_USER_PROMPT_TEMPLATE = """参考资料:
{context}

---

根据上面的参考资料，回答以下问题。如果参考资料中没有直接答案，请结合你的医学护理专业知识进行回答。请确保回答专业、准确、全面。

问题: {question}
"""

# 定义命令行参数解析函数
def parse_args():
    # 创建参数解析器实例
    parser = argparse.ArgumentParser(description="批量推理脚本")
    # 添加输入文件路径参数
    parser.add_argument("--input_file", type=str, required=True, help="输入 JSONL 文件路径")
    # 添加输出文件路径参数
    parser.add_argument("--output_file", type=str, required=True, help="输出 JSONL 文件路径")
    # 添加模型路径参数
    parser.add_argument("--model_path", type=str, required=True, help="模型路径路径")
    # 添加GPU数量参数
    parser.add_argument("--gpus", type=int, default=1, help="使用的 GPU 数量")
    # 新增：批量推理的批次大小
    parser.add_argument("--batch_size", type=int, default=32, help="每次推理的批次大小")
    # 新增：VLLM使用的GPU内存比例
    parser.add_argument("--gpu_memory_utilization", type=float, default=0.9, help="VLLM使用的GPU内存比例。建议在0.8-0.98之间")
    # 新增：VLLM批处理中的最大token数
    parser.add_argument("--max_num_batched_tokens", type=int, default=8192, help="VLLM批处理中的最大token数。根据模型上下文窗口和batch_size调整")
    
    # 返回解析后的参数
    return parser.parse_args()

# --- RAG Setup Function ---
def setup_rag(data_path: str, embedding_model_name: str, chunk_size: int, chunk_overlap: int) -> Chroma.as_retriever:
    """
    加载数据，分块，生成嵌入，构建向量存储并返回检索器。
    """
    if not os.path.exists(data_path):
        print(f"错误：RAG数据目录 {data_path} 不存在。将仅使用模型自身知识回答。")
        return None
        
    print(f"开始加载和处理RAG数据来自: {data_path}")
    
    try:
        # 1. 加载文档
        loader = DirectoryLoader(data_path, glob="*.txt", loader_cls=TextLoader, show_progress=True)
        documents = loader.load()
        
        if not documents:
            print(f"警告：在 {data_path} 目录中未找到任何 .txt 文件。将仅使用模型自身知识回答。")
            return None
            
        print(f"加载 {len(documents)} 个文档。")

        # 2. 分割文档
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        chunks = text_splitter.split_documents(documents)
        print(f"分割成 {len(chunks)} 个文本块。")

        # 3. 初始化嵌入模型
        print(f"初始化嵌入模型: {embedding_model_name}")
        # model_kwargs={'device': 'cuda'} 可以指定GPU，但sentence-transformers默认会尝试GPU
        # encode_kwargs={'normalize_embeddings': True} 是一个常见选项
        embeddings = HuggingFaceBgeEmbeddings(model_name=embedding_model_name, model_kwargs={'device': 'cuda'})
        print("嵌入模型初始化完成。")

        # 4. 构建向量存储
        print("开始构建向量存储 (Chroma)...")
        # 可以选择持久化到磁盘，这里为了简单，使用内存存储
        # persist_directory = "./chroma_db"
        # if os.path.exists(persist_directory):
        #     print(f"加载现有向量存储: {persist_directory}")
        #     vectorstore = Chroma(persist_directory=persist_directory, embedding_function=embeddings)
        # else:
        #     print(f"创建新的向量存储并保存到: {persist_directory}")
        #     vectorstore = Chroma.from_documents(chunks, embeddings, persist_directory=persist_directory)
        #     vectorstore.persist() # 如果持久化，需要调用persist()

        # 使用内存存储
        vectorstore = Chroma.from_documents(chunks, embeddings)
        print("向量存储构建完成。")

        # 5. 创建检索器
        retriever = vectorstore.as_retriever(search_kwargs={"k": RAG_SEARCH_K})
        print("RAG 设置完成，检索器已就绪。")
        return retriever

    except Exception as e:
        print(f"RAG setup 过程中发生错误: {e}")
        print("将仅使用模型自身知识回答。")
        return None


def main():
    # 解析命令行参数
    args = parse_args()

    # --- RAG Setup ---
    retriever = setup_rag(RAG_DATA_PATH, RAG_EMBEDDING_MODEL_NAME, RAG_CHUNK_SIZE, RAG_CHUNK_OVERLAP)

    # 加载分词器
    print(f"加载tokenizer来自: {args.model_path}")
    tokenizer = AutoTokenizer.from_pretrained(args.model_path, trust_remote_code=True)
    print("Tokenizer加载完成。")

    # 初始化LLM模型，设置GPU并行数量和VLLM优化参数
    print(f"初始化LLM模型来自: {args.model_path}")
    llm = LLM(
        model=args.model_path,
        tensor_parallel_size=args.gpus,
        gpu_memory_utilization=args.gpu_memory_utilization,
        max_num_batched_tokens=args.max_num_batched_tokens,
        trust_remote_code=True # 需要加上，Qwen模型需要
    )
    print("LLM模型初始化完成。")

    # 设置生成参数
    sampling_params = SamplingParams(
        temperature=0.0,    # 降低温度，鼓励模型生成更确定、更少随机性的回答，这有助于提高准确性
        top_p=0.9,          # 稍微收紧top_p，进一步减少生成的多样性
        top_k=5,            # 降低top_k，进一步收敛生成范围
        max_tokens=2048,    # 根据实际答案长度调整，避免截断或冗余。
        # 如果模型有特定的停止标记，例如 Qwen 的 <|im_end|>，可以添加
        # stop=["<|im_end|>"] # Qwen模型的停止标记
    )
    # 注意：VLLM 的 stop 参数是字符串列表，不是 token ID 列表

    # 确保输出文件目录存在
    output_dir = os.path.dirname(args.output_file)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    # 读取输入文件
    input_data = []
    with open(args.input_file, 'r', encoding='utf-8') as f:
        for line in f:
            if line.strip():
                try:
                    input_data.append(json.loads(line))
                except json.JSONDecodeError:
                    print(f"Warning: 跳过输入文件中无效的JSON行: {line.strip()}")
    
    if not input_data:
        print("输入文件中没有找到任何有效的问答对。")
        return

    print(f"载入 {len(input_data)} 条问答对进行推理。")

    # 实现断点续传
    processed_ids = set()
    existing_outputs = []
    if os.path.exists(args.output_file):
        print(f"检测到输出文件 {args.output_file}，加载已处理的ID...")
        try:
            with open(args.output_file, 'r', encoding='utf-8') as f_out_exist:
                for line in f_out_exist:
                    try:
                        item = json.loads(line)
                        if "id" in item:
                            processed_ids.add(item["id"])
                            existing_outputs.append(item) # 保留现有结果，最后一起写回（可选，直接追加更简单）
                    except json.JSONDecodeError:
                        pass # 忽略无效行
            print(f"已加载 {len(processed_ids)} 条已处理的问答对ID。")
        except Exception as e:
            print(f"加载现有输出文件出错: {e}. 将忽略现有内容并从头开始处理未完成的样本。")
            processed_ids = set()
            existing_outputs = []


    # 过滤掉已处理的样本
    items_to_process = [item for item in input_data if item["id"] not in processed_ids]
    # 按照ID排序，以确保结果顺序和断点续传的连续性 (如果原始文件不是按ID排序的话)
    # items_to_process.sort(key=lambda x: x["id"]) # 假设输入文件已排序或ID可以作为有效的排序键

    if not items_to_process:
        print("所有问答对都已处理。")
        # 统计最终输出文件中的总行数
        final_output_count = 0
        if os.path.exists(args.output_file):
            with open(args.output_file, 'r', encoding='utf-8') as f:
                for _ in f:
                    final_output_count += 1
        print(f"最终输出文件中包含 {final_output_count} 条结果。")
        return

    print(f"剩余 {len(items_to_process)} 条问答对待处理。")

    # 存储所有Prompt
    all_prompts = []
    # 存储对应的原始item，用于结果输出
    original_items_for_batch = []

    # --- 准备带RAG上下文的Prompt ---
    print("准备带RAG上下文的Prompt...")
    for item in tqdm(items_to_process, desc="准备Prompt"):
        question = item.get("question", "")
        
        # 1. RAG 检索
        context = ""
        if retriever:
            try:
                retrieved_docs: List[Document] = retriever.get_relevant_documents(question)
                # 将检索到的文档内容拼接起来作为上下文
                context = "\n---\n".join([doc.page_content for doc in retrieved_docs])
            except Exception as e:
                 print(f"Warning: 对问题 ID={item.get('id')} 进行RAG检索时发生错误: {e}. 跳过RAG，仅使用模型知识。")
                 context = "" # 清空上下文，退化为非RAG模式

        # 2. 构建 RAG 增强的用户 Prompt
        # 使用模板，如果context为空，模板中的 {context} 会是空的，但Prompt结构还在
        

        user_prompt_content = RAG_USER_PROMPT_TEMPLATE.format(context=context, question=question)
        
        # user_prompt = (
        # "你是一名专业的医疗护理专家。请根据你的专业知识，准确、详细地回答以下护理问题。\n"
        # f"【问题】{user_prompt_content}"
        # )

        # 3. 构建消息格式，加入System Prompt
        messages = [
            {"role": "user", "content": user_prompt_content}
        ]
        
        # 使用Qwen的chat template生成最终Prompt
        # 需要确保 tokenizer 已经正确加载并且支持 apply_chat_template
        try:
            prompt = tokenizer.apply_chat_template(
                messages,
                tokenize=False,           # 不进行分词
                add_generation_prompt=True,  # 添加生成提示 (例如 Qwen 的 <|im_start|>assistant\n)
                enable_thinking=False     # Qwen可能没有这个参数，根据模型实际支持来
            )
        except Exception as e:
             print(f"Error applying chat template for item ID={item.get('id')}: {e}")
             # 如果模板应用失败，可以尝试一个简单的 fallback 结构
             prompt = f"{RAG_USER_PROMPT_TEMPLATE}\n{user_prompt_content}"

        all_prompts.append(prompt)
        original_items_for_batch.append(item) # 存储对应的原始item

    # 以追加模式写入结果，确保断点续传时不会覆盖
    with open(args.output_file, 'a', encoding='utf-8') as f_out:
        # 批量处理所有Prompt
        num_batches = math.ceil(len(all_prompts) / args.batch_size)
        
        # tqdm 用于显示批次处理进度
        for i in tqdm(range(num_batches), desc="批量推理进度"):
            batch_start = i * args.batch_size
            batch_end = min((i + 1) * args.batch_size, len(all_prompts))
            
            current_batch_prompts = all_prompts[batch_start:batch_end]
            # 获取当前批次对应的原始 item 对象
            current_batch_original_items = original_items_for_batch[batch_start:batch_end]

            if not current_batch_prompts:
                continue # 跳过空批次

            try:
                # 批量生成回答
                results = llm.generate(current_batch_prompts, sampling_params)
                
                # result 是一个 RequestOutput 对象列表，与 current_batch_prompts 一一对应
                for j, result in enumerate(results):
                    original_item = current_batch_original_items[j]
                    
                    # VLLM 生成结果可能有多项 (best_of > 1)，这里取第一个
                    if result.outputs:
                        generated_text = result.outputs[0].text.strip()
                    else:
                         generated_text = "" # 没有生成结果

                    output_item = {
                        "id": original_item.get("id"),
                        "question": original_item.get("question", ""),
                        "answer": original_item.get("answer", ""), # 原始答案也保留，方便eval.py
                        "model_response": generated_text
                    }
                    f_out.write(json.dumps(output_item, ensure_ascii=False) + '\n')
                    f_out.flush() # 立即写入文件，增强断点续传的可靠性

            except Exception as e:
                # 打印错误信息，并尝试记录是哪个批次出的问题
                print(f"\nError during VLLM batch generation (batch {i+1}/{num_batches}, IDs from {current_batch_original_items[0].get('id')}): {e}")
                # 如果发生错误，这批次的数据可能都没写进去。
                # 断点续传会在下次运行时重新处理这些ID。
                # 为了防止数据丢失，我们可能需要在异常处理中也尝试写入一些失败标记，
                # 但对于简单的断点续传，忽略并让下次重试是更常见的做法。
                # 这里只打印警告并继续下一个批次（如果可能）。
                
    print(f"推理完成。所有结果已保存到 {args.output_file}。")
    # 统计最终输出文件中的总行数（包括之前已存在的和本次新生成的）
    final_output_count = 0
    if os.path.exists(args.output_file):
        with open(args.output_file, 'r', encoding='utf-8') as f:
            for _ in f:
                final_output_count += 1
    print(f"最终输出文件中包含 {final_output_count} 条结果。")


# 程序入口点
if __name__ == "__main__":
    main()


rl_trainer.py:
from trl import GRPOConfig, GRPOTrainer
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM
from datasets import load_dataset
from bert_score import BERTScorer
import re
import os
os.environ["TRANSFORMERS_OFFLINE"] = "1"
os.environ["HF_DATASETS_OFFLINE"] = "1"
os.environ["BERTSCORE_RESOURCE_DIR"] = "./bert-zh-model"

# 系统提示调整建议（符合临床步骤）
cot_system_prompt = """你是一个护理学专家，请按以下步骤回答问题：
1. 临床场景分析：识别患者的关键症状和护理需求
2. 理论依据：关联相关护理学理论和操作规范
3. 推理过程：用<reasoning>...</reasoning>包裹逻辑推理
4. 最终答案：用<answer>...</answer>包裹明确结论"""

# 强化学习奖励计算器（适配GRPOTrainer接口）
class CotRewardFunction:
    __name__ = "CotRewardFunction"  # 添加类属性
    def __init__(self, references):
        self.scorer = BERTScorer(
            model_type="/inspire/hdd/project/socialscience/xialingying041-summer-041/project/bert-zh-model",  # 直接使用本地路径作为 model_type
            num_layers=12,                  # bert-base-chinese 的层数
        )
        self.reference_dict = {ref["question"]: ref["answer"] for ref in references}

    def __call__(self, prompts, completions, **kwargs):
        rewards = []
        for prompt, completion in zip(prompts, completions):
            # 格式奖励（30%）
            format_score = self._calc_format_score(completion)
            
            # 内容奖励（70%）
            content_score = self._calc_content_score(prompt, completion)
            
            rewards.append(0.3*format_score + 0.7*content_score)
        return rewards

    def _calc_format_score(self, text):
        """严格格式验证"""
        reasoning_blocks = re.findall(r'<reasoning>\s*(.*?)\s*</reasoning>', text, re.DOTALL)
        answer_blocks = re.findall(r'<answer>\s*(.*?)\s*</answer>', text, re.DOTALL)
        
        if len(reasoning_blocks) !=1 or len(answer_blocks)!=1:
            return 0.0  # 格式错误惩罚
        
        answer_len = len(answer_blocks[0].strip())
        reasoning_len = len(reasoning_blocks[0].strip())
        if reasoning_len >= 30 and answer_len >= 50:
            return 1.0  # 完全符合要求
        elif reasoning_len >= 15 or answer_len >= 25:  # 部分符合
            return 0.5
        else:
            return 0.0  # 即使格式正确但内容过短

    def _find_reference(self, prompt):
        """增强型答案匹配"""
        # 双重匹配策略
        question_from_prompt = re.findall(
            r"(问题|Question)[：:]?\s*(.*?)(?=\n|$)", 
            prompt, 
            re.DOTALL
        )
        # 获取最后一个匹配的问题（假设最新问题在最后）
        current_question = question_from_prompt[-1][1].strip() if question_from_prompt else ""
        
        # 分层查找策略
        reference_answer = self.reference_dict.get(current_question, None)
        if not reference_answer:
            # 模糊匹配：去除标点后的部分匹配
            clean_question = re.sub(r"[^\w\u4e00-\u9fff]", "", current_question)
            for q, a in self.reference_dict.items():
                if clean_question in re.sub(r"[^\w\u4e00-\u9fff]", "", q):
                    return a
            # 最终保底返回
            return "[未找到匹配答案]"
        return reference_answer

    def _calc_content_score(self, prompt, completion):
        """增强的内容评分"""
        gen_answer = re.findall(r"<answer>(.*?)</answer>", completion, re.DOTALL)
        if not gen_answer:
            return 0.0
        
        ref_answer = self._find_reference(prompt)
        if not ref_answer.strip():
            print(f"警告：未找到问题对应的参考答案 | Prompt: {prompt[:100]}...")
            return 0.0
        
        # 过滤空内容
        gen_text = gen_answer[0].strip()
        if not gen_text:
            return 0.0
        
        # 执行评分
        try:
            _, _, F1 = self.scorer.score([gen_text], [ref_answer])
            return F1.mean().item()
        except Exception as e:
            print(f"评分错误：{str(e)}")
            return 0.0


tokenizer = AutoTokenizer.from_pretrained("/inspire/hdd/project/socialscience/xialingying041-summer-041/textbooks/model_results/1k/weights/checkpoint-285")

# 数据预处理
def format_prompt(example):
    messages = [
        {"role": "system", "content": cot_system_prompt},
        {"role": "user", "content": example["question"]}
    ]
    return {"prompt": tokenizer.apply_chat_template(messages, tokenize=False)}

# 主训练流程
def train_rl_cot():
    # 加载微调后的模型
    model = AutoModelForCausalLM.from_pretrained(
        "/inspire/hdd/project/socialscience/xialingying041-summer-041/textbooks/model_results/1k/weights/checkpoint-285",
        # torch_dtype=torch.bfloat16,
        device_map="auto",
        use_cache=False
    )
    tokenizer = AutoTokenizer.from_pretrained("/inspire/hdd/project/socialscience/xialingying041-summer-041/textbooks/model_results/1k/weights/checkpoint-285")

    # 加载数据集
    dataset = load_dataset("json", data_files={"train": "/inspire/hdd/project/socialscience/xialingying041-summer-041/strategyE+SFT+GRPO+pdf+textbook/data.jsonl"})["train"]
    dataset = dataset.map(format_prompt)
    
    # 初始化奖励函数
    reward_func = CotRewardFunction(dataset) 

    
    grpo_config = GRPOConfig(
        output_dir="./rl_results",
        logging_dir="logs",
        per_device_train_batch_size=4,
        gradient_accumulation_steps=8,
        gradient_checkpointing=True,
        fp16=True,  
        bf16=False,
        optim="adafactor",
        learning_rate=1e-7,  
        beta=0.2,     
        lr_scheduler_type="cosine",    
        num_generations=4,  
        max_completion_length=256,  
        temperature=0.4,
        top_p=0.9,
        loss_type="bnpo",   
        num_iterations=2,   
        epsilon=0.2,        
        epsilon_high=0.28,  
        scale_rewards=False, 
        logging_steps=10,
        save_steps=500
    )
    
    torch.cuda.empty_cache()
    torch.backends.cuda.max_split_size_mb = 128

   
    trainer = GRPOTrainer(
        model=model,
        args=grpo_config,
        reward_funcs=reward_func,  
        train_dataset=dataset,
        processing_class=tokenizer,
    )

    # 开始训练
    trainer.train()
    trainer.save_model("./rl_cot_model")

if __name__ == "__main__":
    train_rl_cot()

SFT_trainer.py:
from transformers import AutoTokenizer, AutoModelForCausalLM, TrainingArguments
import torch
from datasets import load_dataset
import os
from trl import SFTTrainer, SFTConfig

os.environ["CUDA_VISIBLE_DEVICES"] = "0"

# 加载模型和分词器
tokenizer = AutoTokenizer.from_pretrained(
    '/inspire/hdd/project/socialscience/public/models/Qwen3-8B',
    trust_remote_code=True,
    padding_side="right"
)
model = AutoModelForCausalLM.from_pretrained(
    '/inspire/hdd/project/socialscience/public/models/Qwen3-8B',
    device_map='auto',
    torch_dtype=torch.bfloat16,
    use_cache=False  # 模型加载时禁用缓存
)

# 确保pad_token配置正确
if tokenizer.pad_token is None:
    tokenizer.pad_token = tokenizer.eos_token
model.config.pad_token_id = tokenizer.pad_token_id


def build_prompt_strategy_e(question, tokenizer):
    """策略 E：结合 CoT 与专家角色设定 (强化版 CoT)"""
    user_prompt = (
        "你是一名资深的医疗护理专家。请你针对以下护理问题，首先进行详细的专业思考，包括但不限于问题的关键点、可能涉及的护理原则、相关医学知识等。然后，基于你的思考，给出全面、准确、专业的最终回答。\n"
        f"【问题】{question}\n"
        "【专业思考过程】"
    )
    messages = [{"role": "user", "content": user_prompt}]
    text = tokenizer.apply_chat_template(
        messages,
        tokenize=False,
        add_generation_prompt=True,
        enable_thinking=True, # 如果分词器支持该参数
    )
    return text

def create_prompt(example):
    text = build_prompt_strategy_e(example["question"], tokenizer)
    example["prompt"] = text
    return example

def process_func(example):
    MAX_LENGTH = 1024
    instruction = tokenizer(example["prompt"], 
                          add_special_tokens=False,
                          truncation=True,
                          max_length=MAX_LENGTH)
    
    response = tokenizer(f"{example['answer']}<|eot_id|>",
                       add_special_tokens=False,
                       truncation=True,
                       max_length=MAX_LENGTH)
    
    input_ids = instruction["input_ids"] + response["input_ids"] + [tokenizer.pad_token_id]
    attention_mask = instruction["attention_mask"] + response["attention_mask"] + [1]
    labels = [-100]*len(instruction["input_ids"]) + response["input_ids"] + [tokenizer.pad_token_id]
    
    input_ids = input_ids[:MAX_LENGTH]
    attention_mask = attention_mask[:MAX_LENGTH]
    labels = labels[:MAX_LENGTH]
    
    return {
        "input_ids": input_ids,
        "attention_mask": attention_mask,
        "labels": labels
    }

print("-----Load SFT dataset-----")
dataset = load_dataset(
    "json",
    data_files={
        "train": "/inspire/hdd/project/socialscience/xialingying041-summer-041/project/data/new_train_data/combined_data.jsonl"
    }
)["train"]

print("-----train_test_split-----")
actual_sample_size = len(dataset)
sampled_dataset = dataset.shuffle(seed=42).select(range(actual_sample_size))
train_test_split = sampled_dataset.train_test_split(test_size=0.1, seed=42)
train_dataset = train_test_split["train"]
eval_dataset = train_test_split["test"]

print("-----create prompt-----")
train_dataset = train_dataset.map(create_prompt)
eval_dataset = eval_dataset.map(create_prompt)

print("-----transfer to tokenized id-----")
tokenized_train = train_dataset.map(
    process_func, 
    remove_columns=train_dataset.column_names,
    batched=False
)
tokenized_eval = eval_dataset.map(
    process_func,
    remove_columns=eval_dataset.column_names,
    batched=False
)

# 修正后的训练参数
args = SFTConfig(
    output_dir="./weights",
    per_device_train_batch_size=2,
    gradient_accumulation_steps=4,
    logging_dir="./logs",
    logging_steps=20,
    eval_steps=30,
    num_train_epochs=10,
    save_steps=300,
    learning_rate=1e-5,
    warmup_ratio=0.1,
    weight_decay=0.01,
    lr_scheduler_type="cosine",
    gradient_checkpointing=False,
    seed=42,
    # bf16=True,
    max_length=1024,
    dataset_text_field="text",
    padding_free=False
)

# 关键修改点：使用 processing_class 传递分词器
trainer = SFTTrainer(
    model=model,
    args=args,
    train_dataset=tokenized_train,
    eval_dataset=tokenized_eval,
    processing_class=tokenizer  
)

print("-----Starting Training-----")
trainer.train()
trainer.save_model("./final_model")
tokenizer.save_pretrained("./final_model")
print("-----SFT succeeded!-----")


README.md:
大模型护理问题应答能力构建课题项目的代码实现

创建者：毛宣杰、许圣瑞、潘亚宁-009

本项目的pipeline包括Caring-QA数据集的搭建、SFT微调与GRPO强化学习训练。如果您想检测项目推理结果，请在submit根目录下运行

**bash inference.sh**

如果您想评估推理结果，请运行

**bash eval.sh**


###数据
我们使用的护理语料数据来自于两部分：项目提供的文献资料库（pdf）和收集的护理专业教科书（textbook）

1.对于pdf原始数据，使用data/pdf下的extract_pdf_text.py将其提取成txt格式，然后使用generate_qa_from_text.py对每篇文章生成10个专业的qa对。

2.对于textbook电子版教材，pdf格式使用data/textbook/extract_pdf_text.py将电子教材提取成txt格式，doc格式则使用然后使用data/textbook/word2txt提取成txt格式，然后使用generate_qa.py对每个分块生成QA对。

使用combine.py将处理后的两部分数据合并，再使用rejection_sampling.py做拒绝采样，过滤后数据保存在data/clean_data下


###第一阶段训练——SFT微调
运行：
python SFT_trainer.py
可以监督微调本地的Qwen-8B模型，微调后的模型保存在data/SFT/SFT_model


###第二阶段训练——GRPO强化学习训练
运行：
python rl_train.py
在SFT微调模型基础上进行强化学习训练，训练后的模型保存在data/model


###RAG推理
执行
python RAG.py
可以通过RAG方式进行推理



extract_pdf_text.py:
import os
from pdfminer.high_level import extract_text
from tqdm import tqdm
import argparse


def extract_text_from_pdf(pdf_path):
    """
    从PDF文件中提取文本内容。
    """
    try:
        text = extract_text(pdf_path)
        return text
    except Exception as e:
        print(f"Error extracting text from {pdf_path}: {e}")
        return None

def process_all_pdfs(input_dir, output_dir):
    """
    处理指定目录下的所有PDF文件，并保存提取的文本。
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    pdf_files = [f for f in os.listdir(input_dir) if f.lower().endswith('.pdf')]

    # 记录已处理的文件数量，用于检查断点续传
    processed_count = 0
    if os.path.exists(os.path.join(output_dir, "processed_count.txt")):
        with open(os.path.join(output_dir, "processed_count.txt"), 'r') as f:
            processed_count = int(f.read().strip())

    # 从上一次处理停止的地方开始
    start_index = processed_count

    print(f"Starting PDF processing from index {start_index}...")

    for i, pdf_file in enumerate(tqdm(pdf_files[start_index:], desc="Processing PDFs")):
        pdf_path = os.path.join(input_dir, pdf_file)
        txt_filename = os.path.splitext(pdf_file)[0] + ".txt"
        output_txt_path = os.path.join(output_dir, txt_filename)

        if os.path.exists(output_txt_path):
            # 如果输出文件已存在，跳过处理，继续下一个
            print(f"Skipping {pdf_file}, text already extracted.")
            continue

        text = extract_text_from_pdf(pdf_path)
        if text:
            with open(output_txt_path, 'w', encoding='utf-8') as f:
                f.write(text)

        # 更新已处理文件数量
        processed_count = start_index + i + 1
        with open(os.path.join(output_dir, "processed_count.txt"), 'w') as f:
            f.write(str(processed_count))

    print(f"Finished processing {len(pdf_files)} PDFs.")

if __name__ == "__main__":
    # parser = argparse.ArgumentParser()
    # parser.add_argument('--input_pdf_dir', type=str, default="/inspire/hdd/project/socialscience/public/references")
    # parser.add_argument('--output_text_dir', type=str, default="/inspire/hdd/project/socialscience/xialingying041-summer-041/pdf/output/extracted_texts")
    # args = parser.parse_args()
    # # input_pdf_dir = "/inspire/hdd/project/socialscience/public/references"  # 项目提供的PDF资料库路径 [cite: 5]
    # # input_pdf_dir = "/inspire/hdd/project/socialscience/xialingying041-summer-041/pdf/re1"  # 项目提供的PDF资料库路径 [cite: 5]
    # # output_text_dir = "/inspire/hdd/project/socialscience/xialingying041-summer-041/pdf/output/extracted_texts"  # 提取文本的输出目录
    
    parser = argparse.ArgumentParser()
    parser.add_argument('--input_pdf_dir', type=str, default="/inspire/hdd/project/socialscience/xialingying041-summer-041/textbooks/pdf1")
    parser.add_argument('--output_text_dir', type=str, default="/inspire/hdd/project/socialscience/xialingying041-summer-041/pdf/output/extracted_textbooks1")
    args = parser.parse_args()

    print(f"Input PDF directory: {args.input_pdf_dir}")
    print(f"Output text directory: {args.output_text_dir}")

    process_all_pdfs(args.input_pdf_dir, args.output_text_dir)
    print("PDF文本提取完成。")


generate_qa_from_text.py:

'''
python xialingying041-summer-041/pdf/generate_qa_from_text.py \
  --input_text_dir /inspire/hdd/project/socialscience/xialingying041-summer-041/pdf/output/extracted_texts \
  --output_jsonl_file /inspire/hdd/project/socialscience/xialingying041-summer-041/pdf/output/generated_train.jsonl \
  --model_path /inspire/hdd/project/socialscience/public/models/Qwen3-32B \
  --gpus 2 \
  --batch_size 8 \
  --num_qa_per_doc 10 \
  --max_retries 1  \
  --start_idx 701 \
  --end_idx 1500
  
--input_text_dir：指向你提取 PDF 文本的目录。
--output_jsonl_file：指定输出的 train.jsonl 文件路径，例如 ./data/generated_train.jsonl。
--model_path：指定用于数据构造的大模型路径。你可以使用 Qwen3-8B，或者根据项目要求选择 Qwen2.5 系列、Deepseek-R1-Distill-Qwen 系列等。
--gpus：根据你可用的 GPU 数量进行设置。
--batch_size 1：这里设置为1，因为我们期望模型每次生成一个文档的多个 QA 对，这通常会产生较长的输出，一次处理一个文档更稳妥。
--num_qa_per_doc 10：明确要求每个文档生成10个 QA 对。
  '''
  
import argparse
import json
import os
from tqdm import tqdm
from transformers import AutoTokenizer
from vllm import LLM, SamplingParams
import time

# 定义数据构造的Prompt模板
# 优化后的Prompt，明确要求生成10个问题
QA_GENERATION_PROMPT = """
你是一名专业的护理知识问答出题人。你的任务是根据给定的文章内容，生成 **10个** 高难度、专业性强、贴近临床实操或考试场景的护理问题，以及每个问题对应的直接、准确、全面的答案。

问题应关注：
1. 护理操作的具体步骤或注意事项。
2. 疾病的评估、诊断、治疗或护理干预措施。
3. 药物的使用、副作用或管理。
4. 特定人群（如孕妇、老年人、儿童）的护理特点。
5. 紧急情况下的应对策略。
6. 考试中可能出现的知识点，例如定义、分类、原理或鉴别。

答案必须：
1. 直接从文章中提取或高度概括，确保真实性和权威性。
2. 详细、准确，足以解决问题。
3. 如果文章内容不足以完全回答问题，请在答案中明确指出，但仍尝试给出文章中能提供的信息。

请确保所有问题和答案都使用清晰、专业的中文。

文章内容：
{text_content}

请严格遵循以下JSON数组格式输出结果，不要包含任何额外信息或解释：
[
  {{
    "question": "在这里填写生成的护理问题1",
    "answer": "在这里填写生成的准确答案1"
  }},
  {{
    "question": "在这里填写生成的护理问题2",
    "answer": "在这里填写生成的准确答案2"
  }},
  // ... 以此类推，直到第10个问答对
  {{
    "question": "在这里填写生成的护理问题10",
    "answer": "在这里填写生成的准确答案10"
  }}
]
"""

def parse_args():
    parser = argparse.ArgumentParser(description="根据文本生成问答对脚本")
    parser.add_argument("--input_text_dir", type=str, required=True, help="包含提取文本的目录")
    parser.add_argument("--output_jsonl_file", type=str, required=True, help="输出 JSONL 文件路径")
    parser.add_argument("--model_path", type=str, required=True, help="用于数据构造的模型路径")
    parser.add_argument("--gpus", type=int, default=1, help="使用的 GPU 数量")
    parser.add_argument("--batch_size", type=int, default=1, help="批量推理的批次大小，每个PDF生成10个问题，建议单次处理一个PDF")
    parser.add_argument("--num_qa_per_doc", type=int, default=10, help="每个文档生成问答对的数量")
    parser.add_argument("--max_retries", type=int, default=3, help="生成失败时的最大重试次数")
    parser.add_argument("--start_idx", type=int, default=0, help="从排序后的PDF文件列表的第几（0-indexed）个文件开始处理")
    parser.add_argument("--end_idx", type=int, default=-1, help="处理到排序后的PDF文件列表的第几（0-indexed）个文件（包含此文件）。-1表示处理到最后一个文件")
    return parser.parse_args()

def main():
    args = parse_args()

    # 加载分词器和模型
    tokenizer = AutoTokenizer.from_pretrained(args.model_path)
    llm = LLM(model=args.model_path, tensor_parallel_size=args.gpus)

    # 设置生成参数
    sampling_params = SamplingParams(
        temperature=0.8,
        top_p=0.9,
        top_k=50,
        max_tokens=2048,      # 增加 max_tokens 以适应生成多个QA对的需求
        stop=["\n]"] # 尝试更精准的停止条件，例如在JSON数组闭合的换行符和右括号处停止
    )

    extracted_text_files_list = []
    # 按照文件名排序，确保每次运行的顺序一致
    text_filenames = sorted([f for f in os.listdir(args.input_text_dir) if f.endswith(".txt")])

    # 根据start_idx和end_idx筛选文件
    start_index = args.start_idx
    end_index = args.end_idx if args.end_idx != -1 else len(text_filenames) - 1 # -1表示到最后一个文件

    if start_index < 0 or start_index >= len(text_filenames):
        print(f"错误：start_idx ({start_index}) 超出文件列表范围 (0 - {len(text_filenames) - 1})。")
        return
    if end_index < start_index or end_index >= len(text_filenames):
        print(f"警告：end_idx ({end_index}) 超出文件列表范围或小于start_idx。将处理到最后一个文件 ({len(text_filenames) - 1})。")
        end_index = len(text_filenames) - 1

    files_to_process = text_filenames[start_index : end_index + 1]

    for filename in files_to_process:
        file_path = os.path.join(args.input_text_dir, filename)
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            if content.strip(): # 避免空文件
                extracted_text_files_list.append((filename, content))

    if not extracted_text_files_list:
        print(f"在 {args.input_text_dir} 中指定范围内没有找到任何文本文件，请确保PDF已成功提取或索引范围正确。")
        return

    # 检查输出文件是否存在，如果存在则读取已有的ID，实现断点续传的ID管理
    output_dir = os.path.dirname(args.output_jsonl_file)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    current_id = 0
    if os.path.exists(args.output_jsonl_file):
        print(f"检测到现有输出文件: {args.output_jsonl_file}。尝试加载已有的问答对以确定起始ID...")
        try:
            with open(args.output_jsonl_file, 'r', encoding='utf-8') as f:
                for line in f:
                    try:
                        item = json.loads(line.strip())
                        if "id" in item:
                            current_id = max(current_id, item["id"] + 1)
                    except json.JSONDecodeError:
                        print(f"Warning: 现有文件中存在无效JSON行: {line.strip()}")
        except Exception as e:
            print(f"读取现有文件时发生错误: {e}")
        print(f"已加载 {current_id} 作为新的问答对起始ID。")

    print(f"将处理从索引 {start_index} 到 {end_index} 的 {len(extracted_text_files_list)} 篇文本文章。正在生成问答对...")

    # 以追加模式打开输出文件
    with open(args.output_jsonl_file, 'a', encoding='utf-8') as f_out:
        for doc_idx, (filename, text_content) in enumerate(tqdm(extracted_text_files_list, desc="生成问答对")):
            # 限制每个文本块的大小
            if len(text_content) > 6000: # 根据模型最大上下文窗口调整
                text_content = text_content[:6000] # 简单截断
            
            generated_count_for_doc = 0
            retries = 0

            while generated_count_for_doc < args.num_qa_per_doc and retries < args.max_retries:
                prompt = QA_GENERATION_PROMPT.format(text_content=text_content)
                messages = [{"role": "user", "content": prompt}]
                formatted_prompt = tokenizer.apply_chat_template(
                    messages,
                    tokenize=False,
                    add_generation_prompt=True,
                    enable_thinking=False
                )

                try:
                    results = llm.generate([formatted_prompt], sampling_params)
                    generated_text = results[0].outputs[0].text.strip()
                    
                    # 尝试修复JSON格式
                    # 如果生成内容不是以[开头，则尝试添加
                    if not generated_text.startswith("["):
                        generated_text = "[" + generated_text
                    # 如果生成内容不是以]结尾，则尝试添加
                    if not generated_text.endswith("]"):
                        generated_text += "]"
                    
                    qa_list = json.loads(generated_text)
                    
                    if not isinstance(qa_list, list):
                        print(f"Warning: 模型生成的不是JSON数组，重试。文件: {filename}, 片段: {text_content[:100]}...")
                        retries += 1
                        # time.sleep(5) # 等待一下再重试
                        continue

                    valid_qa_in_batch = 0
                    for qa_pair in qa_list:
                        if "question" in qa_pair and "answer" in qa_pair:
                            output_item = {
                                "id": current_id,
                                "question": qa_pair["question"].strip(),
                                "answer": qa_pair["answer"].strip()
                            }
                            f_out.write(json.dumps(output_item, ensure_ascii=False) + '\n')
                            current_id += 1
                            valid_qa_in_batch += 1
                        else:
                            print(f"Warning: 模型生成的JSON对象缺少'question'或'answer'字段。文件: {filename}, JSON: {qa_pair}")
                    
                    generated_count_for_doc += valid_qa_in_batch
                    
                    if generated_count_for_doc < args.num_qa_per_doc:
                        print(f"文档 '{filename}' 生成了 {valid_qa_in_batch} 个QA，目标 {args.num_qa_per_doc}。正在重试... (已重试 {retries+1} 次)")
                        retries += 1
                        # time.sleep(5)

                except json.JSONDecodeError:
                    print(f"Warning: 模型生成了无效的JSON输出，重试。文件: {filename}, 片段: {text_content[:100]}...\n{generated_text}")
                    retries += 1
                    # time.sleep(5)
                except Exception as e:
                    print(f"Error generating QA for '{filename}': {e}. 重试... (已重试 {retries+1} 次)")
                    retries += 1
                    # time.sleep(5)
            
            if generated_count_for_doc < args.num_qa_per_doc:
                print(f"警告：文档 '{filename}' 未能生成目标数量 {args.num_qa_per_doc} 个问答对，仅生成了 {generated_count_for_doc} 个。")
            else:
                print(f"文档 '{filename}' 成功生成了 {generated_count_for_doc} 个问答对。")

    print(f"数据构造完成。总共生成了 {current_id} 条问答对。")

if __name__ == "__main__":
    main()



# import argparse
# import json
# import os
# from tqdm import tqdm
# from transformers import AutoTokenizer
# from vllm import LLM, SamplingParams




# # 定义数据构造的Prompt模板
# QA_GENERATION_PROMPT = """
# 请根据以下文章内容，生成一个与文章相关的护理问题和一个准确的答案。答案必须直接从文章中提取或总结，并且能够直接回答问题。问题和答案都必须是中文。

# 文章内容：
# {text_content}

# 请以以下JSON格式输出结果：
# {{
#   "question": "在这里填写生成的问题",
#   "answer": "在这里填写生成的答案"
# }}
# """

# def parse_args():
#     parser = argparse.ArgumentParser(description="根据文本生成问答对脚本")
#     parser.add_argument("--input_text_dir", type=str, required=True, help="包含提取文本的目录")
#     parser.add_argument("--output_jsonl_file", type=str, required=True, help="输出 JSONL 文件路径")
#     parser.add_argument("--model_path", type=str, required=True, help="用于数据构造的模型路径")
#     parser.add_argument("--gpus", type=int, default=1, help="使用的 GPU 数量")
#     return parser.parse_args()

# def main():
#     args = parse_args()

#     # 加载分词器和模型
#     tokenizer = AutoTokenizer.from_pretrained(args.model_path)
#     llm = LLM(model=args.model_path, tensor_parallel_size=args.gpus)

#     # 设置生成参数
#     sampling_params = SamplingParams(
#         temperature=0.7,    # 略高一些的温度，鼓励多样性
#         top_p=0.9,
#         top_k=50,
#         max_tokens=512,      # 确保有足够空间生成问题和答案
#         stop=["\n\n"] # 根据模型生成特点调整，避免生成过多内容
#     )

#     extracted_texts = []
#     for filename in os.listdir(args.input_text_dir):
#         if filename.endswith(".txt"):
#             file_path = os.path.join(args.input_text_dir, filename)
#             with open(file_path, 'r', encoding='utf-8') as f:
#                 content = f.read()
#                 if content.strip(): # 避免空文件
#                     extracted_texts.append(content)

#     if not extracted_texts:
#         print(f"在 {args.input_text_dir} 中没有找到任何文本文件，请确保PDF已成功提取。")
#         return

#     generated_qa_pairs = []

#     # 批量处理文本，提高效率
#     batch_prompts = []
#     # 使用一个计数器来给生成的问答对分配ID
#     current_id = 0

#     print(f"Found {len(extracted_texts)} text files. Generating Q&A pairs...")

#     for text_content in tqdm(extracted_texts, desc="Generating Q&A for texts"):
#         # 限制每个文本块的大小，避免超长输入
#         if len(text_content) > 4000: # 示例截断，根据模型上下文窗口调整
#             text_content = text_content[:4000] + "..." 

#         prompt = QA_GENERATION_PROMPT.format(text_content=text_content)
#         messages = [{"role": "user", "content": prompt}]
#         formatted_prompt = tokenizer.apply_chat_template(
#             messages,
#             tokenize=False,
#             add_generation_prompt=True,
#             enable_thinking=False
#         )
#         batch_prompts.append(formatted_prompt)

#         # 每处理一定数量的文本后进行一次批量推理
#         if len(batch_prompts) >= 16: # 批量大小可以根据GPU显存调整
#             results = llm.generate(batch_prompts, sampling_params)
#             for result in results:
#                 generated_text = result.outputs[0].text.strip()
#                 try:
#                     qa_pair = json.loads(generated_text)
#                     if "question" in qa_pair and "answer" in qa_pair:
#                         qa_pair["id"] = current_id
#                         generated_qa_pairs.append(qa_pair)
#                         current_id += 1
#                     else:
#                         print(f"Warning: Model generated invalid JSON format for text: {generated_text}")
#                 except json.JSONDecodeError:
#                     print(f"Warning: Model generated non-JSON output for text: {generated_text}")
#             batch_prompts = [] # 清空批次

#     # 处理剩余的批次
#     if batch_prompts:
#         results = llm.generate(batch_prompts, sampling_params)
#         for result in results:
#             generated_text = result.outputs[0].text.strip()
#             try:
#                 qa_pair = json.loads(generated_text)
#                 if "question" in qa_pair and "answer" in qa_pair:
#                     qa_pair["id"] = current_id
#                     generated_qa_pairs.append(qa_pair)
#                     current_id += 1
#                 else:
#                     print(f"Warning: Model generated invalid JSON format for text: {generated_text}")
#             except json.JSONDecodeError:
#                 print(f"Warning: Model generated non-JSON output for text: {generated_text}")

#     # 将结果写入输出文件
#     output_dir = os.path.dirname(args.output_jsonl_file)
#     if output_dir:
#         os.makedirs(output_dir, exist_ok=True)

#     with open(args.output_jsonl_file, 'w', encoding='utf-8') as f:
#         for item in generated_qa_pairs:
#             f.write(json.dumps(item, ensure_ascii=False) + '\n')

#     print(f"成功生成 {len(generated_qa_pairs)} 条问答对并保存到 {args.output_jsonl_file}")

# if __name__ == "__main__":
#     main()


word2txt.py:
# 批量处理代码
import os
import re
from docx import Document

def extract_text_from_docx(docx_path, txt_path):
    """从单个 .docx 文件提取文本并保存到 .txt 文件，同时删除文字之间的空格，保留换行符"""
    doc = Document(docx_path)
    
    text = []
    for para in doc.paragraphs:
        para_text = para.text.strip()  # 去掉首尾空格
        # 删除文字之间的所有空格
        para_text = re.sub(r'\s+', '', para_text)
        
        if para_text:
            text.append(para_text)
    
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(text))
    
    print(f"文本已成功提取到 {txt_path}")

def process_all_docx_in_folder(folder_path, output_folder, record_file):
    """处理文件夹中的所有 .docx 文件并支持断点续传"""
    # 获取已处理文件的记录
    processed_files = set()
    if os.path.exists(record_file):
        with open(record_file, 'r', encoding='utf-8') as record_f:
            processed_files = set(record_f.read().splitlines())
    
    # 遍历文件夹中的所有文件
    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            docx_path = os.path.join(folder_path, filename)
            txt_filename = filename.replace('.docx', '.txt')
            txt_path = os.path.join(output_folder, txt_filename)
            
            # 跳过已处理的文件
            if filename in processed_files:
                print(f"跳过已处理的文件: {filename}")
                continue
            
            # 提取文本并保存到 .txt 文件
            extract_text_from_docx(docx_path, txt_path)
            
            # 记录已处理的文件
            with open(record_file, 'a', encoding='utf-8') as record_f:
                record_f.write(f"{filename}\n")

# 设置文件夹路径和记录文件
folder_path = './word'  # .docx 文件所在的文件夹
output_folder = './txt'  # 输出的 .txt 文件所在的文件夹
record_file = './word/processed_files.txt'  # 记录已处理文件的文件

# 创建输出文件夹（如果不存在）
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# 处理所有 .docx 文件
process_all_docx_in_folder(folder_path, output_folder, record_file)

# # 单个文件处理代码
# from docx import Document

# def extract_all_text(docx_path, txt_path):
#     # 打开 docx 文件
#     doc = Document(docx_path)
    
#     # 创建一个列表保存所有段落的文本
#     text = []
    
#     # 遍历每个段落
#     for para in doc.paragraphs:
#         para_text = para.text.strip()
        
#         # 忽略空段落
#         if para_text:
#             text.append(para_text)
    
#     # 将提取的文本写入 .txt 文件
#     with open(txt_path, 'w', encoding='utf-8') as f:
#         f.write("\n".join(text))
    
#     print(f"所有文本内容已经成功提取到 {txt_path}")


# docx_path = './word/基础护理学.docx'  # 替换为实际的 .docx 文件路径
# txt_path = './txt/基础护理学.txt'  # 输出的 .txt 文件路径

# extract_all_text(docx_path, txt_path)


combine.py:
import json
import os

def renumber_ids_and_combine_jsonl(folder_path, output_filename="renumbered_combined_data.jsonl"):
    """
    将指定文件夹下的所有JSON Lines文件的数据整合到一个新的JSON Lines文件中，
    并重新编号每个JSON对象的 'id' 字段，使其从1开始连贯计数。

    Args:
        folder_path (str): 包含JSON Lines文件的文件夹路径。
        output_filename (str): 输出的JSON Lines文件名，默认为 "renumbered_combined_data.jsonl"。
    """
    output_file_path = os.path.join(os.getcwd(), output_filename)
    global_index = 0
    total_processed = 0

    try:
        with open(output_file_path, 'w', encoding='utf-8') as outfile:
            for filename in os.listdir(folder_path):
                if filename.endswith(".jsonl"):
                    file_path = os.path.join(folder_path, filename)
                    print(f"正在处理文件: {filename}")
                    try:
                        with open(file_path, 'r', encoding='utf-8') as infile:
                            for line_num, line in enumerate(infile, 1):
                                try:
                                    data = json.loads(line.strip())
                                    # 重新设置 'id' 的值
                                    data['id'] = global_index
                                    json.dump(data, outfile, ensure_ascii=False)
                                    outfile.write('\n')
                                    global_index += 1
                                    total_processed += 1
                                except json.JSONDecodeError:
                                    print(f"警告: 文件 '{filename}' 第 {line_num} 行不是有效的JSON。已跳过此行。")
                                except KeyError:
                                    print(f"警告: 文件 '{filename}' 第 {line_num} 行的JSON不包含 'id' 字段。已跳过此行。")
                                except Exception as e:
                                    print(f"警告: 处理文件 '{filename}' 第 {line_num} 行时发生未知错误: {e}。已跳过此行。")
                    except FileNotFoundError:
                        print(f"错误: JSON Lines文件未找到: {filename}")
                    except Exception as e:
                        print(f"处理文件 {filename} 时发生未知错误: {e}")

        if total_processed > 0:
            print(f"\n已将所有JSON Lines文件的数据整合到: {output_file_path}")
            print(f"并重新编号了 {total_processed} 条数据的 'id'，起始ID为 0。")
        else:
            print("文件夹中没有找到任何JSON Lines文件。")

    except Exception as e:
        print(f"创建或写入输出文件 '{output_file_path}' 时发生错误: {e}")

if __name__ == "__main__":
    folder_path = "/inspire/hdd/project/socialscience/xialingying041-summer-041/project/data/pdf_train_data"
    output_filename = "/inspire/hdd/project/socialscience/xialingying041-summer-041/project/data/pdf_train_data/combined_data.jsonl"

    if not os.path.isdir(folder_path):
        print(f"错误: 提供的文件夹路径 '{folder_path}' 不存在或不是一个有效的文件夹。")
    else:
        renumber_ids_and_combine_jsonl(folder_path, output_filename)
        



rejection_sampling.py:
import argparse
import json
import os
from tqdm import tqdm
from transformers import AutoTokenizer
from vllm import LLM, SamplingParams

# 评分Prompt模板
SCORE_PROMPT = """
以下是一个护理问题及其回答。请根据以下六个标准对该问答进行评分，评分区间为1-10，其中1表示极差，10表示非常优秀。

1. **准确性**（Accuracy）：答案是否准确，是否符合护理实践中的标准和知识？
2. **相关性**（Relevance）：答案是否直接回答了问题，是否提供了解决方案或实用的护理建议？
3. **清晰度**（Clarity）：答案是否表达清晰，易于理解？是否适合患者或护理人员阅读？
4. **自然性**（Naturalness）：答案是否流畅自然，是否显得不机械或过于公式化？
5. **完整性**（Completeness）：答案是否涵盖了问题的所有方面，是否遗漏重要信息？
6. **多样性**（Diversity）：答案是否能够涵盖护理领域的多样性，避免单一或过于狭窄的视角？

请按格式返回评分，格式如下：
{
  "accuracy": 8,
  "relevance": 9,
  "clarity": 7,
  "naturalness": 9,
  "completeness": 8,
  "diversity": 7
}
"""

def parse_args():
    parser = argparse.ArgumentParser(description="评估已生成的护理问答对并打分")
    parser.add_argument("--input_jsonl", type=str, required=True, help="输入的包含问答对的jsonl文件")
    parser.add_argument("--output_dir", type=str, required=True, help="输出评分结果的目录")
    parser.add_argument("--model_path", type=str, required=True, help="用于评分的模型路径")
    parser.add_argument("--gpus", type=int, default=1, help="使用的 GPU 数量")
    return parser.parse_args()

def score_qa_pairs(qa_pairs, llm, tokenizer, sampling_params):
    scored_qa_pairs = []
    for qa_pair in qa_pairs:
        prompt = SCORE_PROMPT + "\n问题：\n" + qa_pair["question"] + "\n答案：\n" + qa_pair["answer"]
        messages = [{"role": "user", "content": prompt}]
        formatted_prompt = tokenizer.apply_chat_template(
            messages,
            tokenize=False,
            add_generation_prompt=True,
            enable_thinking=False
        )

        # 进行模型推理
        result = llm.generate([formatted_prompt], sampling_params)
        scored_qa_pair = result[0].outputs[0].text.strip()
        
        try:
            # 解析返回的评分
            score = json.loads(scored_qa_pair)
            scored_qa_pairs.append({
                "question": qa_pair["question"],
                "answer": qa_pair["answer"],
                "score": score
            })
        except json.JSONDecodeError:
            print(f"Warning: 模型评分生成失败：\n{scored_qa_pair}")

    return scored_qa_pairs

def main():
    args = parse_args()

    # 加载分词器和模型
    tokenizer = AutoTokenizer.from_pretrained(args.model_path)
    llm = LLM(model=args.model_path, tensor_parallel_size=args.gpus)

    # 设置生成参数
    sampling_params = SamplingParams(
        temperature=0.8,    # 提高温度，鼓励生成更复杂的问答
        top_p=0.9,
        top_k=50,
        max_tokens=1024,      # 增加max_tokens以适应更长、更详细的答案
        stop=["}}"] # 严格控制输出格式，在JSON闭合括号处停止
    )

    # 读取输入的jsonl文件
    with open(args.input_jsonl, 'r', encoding='utf-8') as f_in:
        qa_pairs = [json.loads(line.strip()) for line in f_in.readlines()]

    if not qa_pairs:
        print(f"在 {args.input_jsonl} 中没有找到任何问答对。")
        return

    # 检查输出目录是否存在
    if not os.path.exists(args.output_dir):
        os.makedirs(args.output_dir)

    print(f"读取了 {len(qa_pairs)} 个问答对。正在进行评分...")

    # 评分每个生成的问答对
    scored_qa_pairs = score_qa_pairs(qa_pairs, llm, tokenizer, sampling_params)

    # 保存评分结果
    output_jsonl_file = os.path.join(args.output_dir, "scored_qa_pairs.jsonl")
    with open(output_jsonl_file, 'w', encoding='utf-8') as f_out:
        for scored_qa in scored_qa_pairs:
            f_out.write(json.dumps(scored_qa, ensure_ascii=False) + '\n')

    print(f"评分完成，结果已保存到 {output_jsonl_file}")

if __name__ == "__main__":
    main()

